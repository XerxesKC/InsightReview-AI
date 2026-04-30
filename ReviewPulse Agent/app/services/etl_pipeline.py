from __future__ import annotations

import os
import re
import json
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from pypdf import PdfReader
import pandas as pd
from docx import Document
from langchain_text_splitters import (
    CharacterTextSplitter,
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
    TokenTextSplitter
)
from langchain_core.documents import Document as LangchainDocument


def clean_text(text: str) -> str:
    """
    清洗文本中的乱码和特殊字符
    """
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = '\n'.join(line.strip() for line in text.splitlines())
    return text.strip()


class DocumentTextSplitter:
    """
    文档文本拆分器，支持多种拆分策略
    """

    def __init__(
            self,
            chunk_size: int = 500,
            chunk_overlap: int = 50,
            separator: str = "\n\n",
            splitter_type: str = "recursive",  
            keep_separator: bool = False,
            model_name: str = "gpt-3.5-turbo"  
    ):
        """
        初始化拆分器

        参数:
            chunk_size: 每块的最大大小
            chunk_overlap: 块之间的重叠字符数
            separator: 分隔符
            splitter_type: 拆分器类型
            keep_separator: 是否保留分隔符
            model_name: token拆分器使用的模型名
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separator = separator
        self.splitter_type = splitter_type
        self.keep_separator = keep_separator
        self.model_name = model_name
        self.splitter = self._create_splitter()
        self.stats = {}

    def _create_splitter(self):
        """根据类型创建对应的拆分器"""

        if self.splitter_type == "character":
            return CharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separator=self.separator,
                keep_separator=self.keep_separator
            )

        elif self.splitter_type == "recursive":
            return RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=["\n\n", "\n", "。", "；", "，", " ", ""],
                keep_separator=self.keep_separator
            )

        elif self.splitter_type == "markdown":
            headers_to_split_on = [
                ("#", "Header 1"),
                ("##", "Header 2"),
                ("###", "Header 3"),
                ("####", "Header 4"),
            ]
            return MarkdownHeaderTextSplitter(
                headers_to_split_on=headers_to_split_on,
                return_each_line=False
            )

        elif self.splitter_type == "token":
            return TokenTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                model_name=self.model_name
            )

        else:
            raise ValueError(f"不支持的拆分器类型: {self.splitter_type}")

    def split_text(self, text: str) -> List[str]:
        """
        拆分纯文本

        参数:
            text: 要拆分的文本

        返回:
            拆分后的文本块列表
        """
        if not text:
            return []

        if self.splitter_type == "markdown":
            docs = self.splitter.split_text(text)
            chunks = [doc.page_content for doc in docs]
        else:
            chunks = self.splitter.split_text(text)

        self._update_stats(text, chunks)
        return chunks

    def split_documents(self, documents: Union[List[LangchainDocument], LangchainDocument]) -> List[LangchainDocument]:
        """
        拆分Document对象

        参数:
            documents: Document对象或Document列表

        返回:
            拆分后的Document列表
        """
        if isinstance(documents, LangchainDocument):
            documents = [documents]

        all_text = "\n\n".join([doc.page_content for doc in documents])

        split_docs = self.splitter.split_documents(documents)

        chunk_texts = [doc.page_content for doc in split_docs]
        self._update_stats(all_text, chunk_texts)

        return split_docs

    def _update_stats(self, original_text: str, chunks: List[str]):
        """更新拆分统计信息"""
        if not chunks:
            self.stats = {
                "original_length": len(original_text),
                "num_chunks": 0,
                "chunk_sizes": [],
                "avg_chunk_size": 0,
                "min_chunk_size": 0,
                "max_chunk_size": 0,
            }
            return

        self.stats = {
            "original_length": len(original_text),
            "num_chunks": len(chunks),
            "chunk_sizes": [len(chunk) for chunk in chunks],
            "avg_chunk_size": sum(len(chunk) for chunk in chunks) / len(chunks),
            "min_chunk_size": min([len(chunk) for chunk in chunks]),
            "max_chunk_size": max([len(chunk) for chunk in chunks]),
        }


def get_project_root() -> Path:
    """获取项目根目录"""
    return Path(__file__).parent.parent.parent


def get_absolute_path(relative_path: str) -> str:
    """将相对路径转换为绝对路径"""
    if not relative_path:
        return relative_path

    if os.path.isabs(relative_path):
        return relative_path

    project_root = get_project_root()
    absolute_path = str(project_root / relative_path)

    return absolute_path


def extract_text_from_pdf(file_path: str) -> str:
    """
    解析 PDF 文件，保留标题、段落（简单格式）
    """
    text_content = []
    with open(file_path, 'rb') as f:
        reader = PdfReader(f)
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                cleaned = clean_text(page_text)
                if cleaned:
                    text_content.append(f"[Page {page_num + 1}]\n{cleaned}")
    return '\n\n'.join(text_content)


def extract_text_from_docx(file_path: str) -> str:
    """
    解析 DOCX 文件，保留标题、列表、段落结构
    """
    doc = Document(file_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name.lower() if para.style and para.style.name else ''
        if 'title' in style_name or 'heading' in style_name:
            text = f"## {text} ##"  

        if para._element.pPr is not None and para._element.pPr.numPr is not None:
            text = f"• {text}"  

        paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append("[Table Row] " + " | ".join(row_texts))

    full_text = '\n'.join(paragraphs)
    return clean_text(full_text)


def extract_text_from_markdown(file_path: str) -> str:
    """
    解析 Markdown 文件，保留格式标记用于后续拆分
    """
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()

    lines = content.split('\n')
    processed_lines = []

    for line in lines:
        if re.match(r'^#{1,6}[^#\s]', line):
            match = re.match(r'^(#{1,6})(.+)$', line)
            if match:
                line = f"{match.group(1)} {match.group(2).strip()}"
        processed_lines.append(line)

    return '\n'.join(processed_lines)


def extract_text_from_csv(file_path: str, **kwargs) -> str:
    """
    解析 CSV 文件，使用 pandas 读取并转换为文本格式

    参数:
        file_path: CSV文件路径
        **kwargs: 传递给 pandas.read_csv 的参数，如 encoding, sep, delimiter 等
    """
    try:
        read_params = {
            'encoding': 'utf-8',
            'keep_default_na': False,
            'na_values': [''],
            'engine': 'python',  
            'on_bad_lines': 'skip'  
        }
        read_params.update(kwargs)

        df = None
        errors = []

        separators = [',', ';', '\t', '|', ' ']

        for sep in separators:
            for encoding in ['ANSI','utf-8', 'gbk', 'gb2312', 'latin-1', 'cp1252']:
                try:
                    read_params['encoding'] = encoding
                    read_params['sep'] = sep
                    df = pd.read_csv(file_path, **read_params)
                    if len(df.columns) > 1:  
                        break
                except Exception as e:
                    errors.append(f"{encoding}/{sep}: {str(e)}")
                    continue
            if df is not None and len(df.columns) > 1:
                break

        if df is None or len(df.columns) <= 1:
            try:
                df = pd.read_csv(file_path, engine='python', encoding='utf-8',
                                 skipinitialspace=True, on_bad_lines='skip')
            except:
                try:
                    df = pd.read_csv(file_path, header=None, names=['content'],
                                     encoding='utf-8', on_bad_lines='skip')
                except:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        lines = f.readlines()
                    df = pd.DataFrame({'content': [line.strip() for line in lines if line.strip()]})

        text_lines = []

        if df is not None and not df.empty:
            columns = df.columns.tolist()
            text_lines.append("【数据表字段】")
            text_lines.append(" | ".join([f"{col}" for col in columns]))
            text_lines.append("")

            text_lines.append("【数据内容】")

            max_rows = min(100, len(df))
            for idx in range(max_rows):
                row = df.iloc[idx]
                row_text = " | ".join(
                    [f"{col}: {str(row[col])}" for col in columns if pd.notna(row[col]) and str(row[col]).strip()])
                if row_text:
                    text_lines.append(f"第{idx + 1}行: {row_text}")

            if len(df) > max_rows:
                text_lines.append(f"... 共 {len(df)} 行，仅显示前 {max_rows} 行")

            text_lines.append("")
            text_lines.append("【数据统计】")
            text_lines.append(f"总行数: {len(df)}")
            text_lines.append(f"总列数: {len(columns)}")
        else:
            text_lines.append("【CSV文件为空或无法解析】")

        return '\n'.join(text_lines)

    except Exception as e:
        print(f"CSV解析失败: {e}")
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            return f"【CSV解析失败，按文本读取】\n{content}"
        except:
            return "【CSV文件解析失败】"


def extract_text_from_txt(file_path: str) -> str:
    """
    解析 TXT 文件，支持多种编码
    """
    encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1', 'cp1252', 'ascii']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            print(f"使用编码 {encoding} 读取失败: {e}")
            continue

    with open(file_path, 'rb') as f:
        content = f.read().decode('utf-8', errors='ignore')
        return content


def parse_document(file_path: str, **kwargs) -> str:
    """
    根据文件类型自动选择解析器

    参数:
        file_path: 文件路径
        **kwargs: 传递给特定解析器的额外参数（如CSV的sep, encoding等）
    """
    file_path = get_absolute_path(file_path)
    path = Path(file_path)
    suffix = path.suffix.lower()

    print(f"正在解析: {path.name}")

    if suffix == '.pdf':
        raw = extract_text_from_pdf(file_path)
    elif suffix == '.docx':
        raw = extract_text_from_docx(file_path)
    elif suffix == '.csv':
        csv_params = {k.replace('csv_', ''): v for k, v in kwargs.items() if k.startswith('csv_')}
        raw = extract_text_from_csv(file_path, **csv_params)
    elif suffix == '.md':
        raw = extract_text_from_markdown(file_path)
    elif suffix == '.txt':
        raw = extract_text_from_txt(file_path)
    elif suffix == '.json':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            raw = json.dumps(data, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"JSON解析失败: {e}")
            raw = path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(f"不支持的文件格式: {suffix}")

    final_text = clean_text(raw)
    return final_text


def build_document_chunks(
        file_path: str | Path,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        splitter_type: str = "recursive",
        **kwargs
) -> list[dict[str, str | int | dict]]:
    """
    构建文档块，整合了解析和拆分功能

    参数:
        file_path: 文件路径
        chunk_size: 块大小
        chunk_overlap: 块重叠大小
        splitter_type: 拆分器类型
        **kwargs: 传递给 parse_document 的额外参数

    返回:
        文档块列表，每块包含内容、来源、索引和统计信息
    """
    path = Path(file_path)
    if not path.exists():
        print(f"文件不存在: {path}")
        return []

    try:
        content = parse_document(str(path), **kwargs)
    except Exception as e:
        print(f"解析文档失败: {e}")
        return []

    splitter = DocumentTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        splitter_type=splitter_type
    )

    chunks = splitter.split_text(content)

    file_stats = path.stat()
    file_info = {
        "file_name": path.name,
        "file_size": file_stats.st_size,
        "file_type": path.suffix.lower(),
        "modified_time": file_stats.st_mtime
    }

    return [
        {
            "content": chunk,
            "source": path.name,
            "chunk_index": index,
            "file_info": file_info,
            "splitter_stats": splitter.stats
        }
        for index, chunk in enumerate(chunks)
    ]


def example_usage():
    """展示如何使用新增的功能"""

    csv_file = "data/uploads/merchant_1001/川香居菜单.csv"
    csv_file = get_absolute_path(csv_file)
    csv_chunks = build_document_chunks(
        csv_file,
        chunk_size=300,
        splitter_type="recursive",
        csv_encoding='utf-8',  
        csv_sep=','  
    )
    print(f"CSV文件生成 {len(csv_chunks)} 个文本块")

    md_file = "data/uploads/merchant_1001/品牌故事与主厨介绍.md"
    md_file = get_absolute_path(md_file)
    md_chunks = build_document_chunks(
        md_file,
        chunk_size=400,
        splitter_type="markdown"  
    )
    print(f"Markdown文件生成 {len(md_chunks)} 个文本块")

    txt_file = "data/uploads/merchant_1001/顾客常见问题解答 .txt"
    txt_file = get_absolute_path(txt_file)
    txt_chunks = build_document_chunks(
        txt_file,
        chunk_size=200,
        chunk_overlap=30
    )
    print(f"TXT文件生成 {len(txt_chunks)} 个文本块")

    if csv_chunks:
        print(f"CSV第一个块: {csv_chunks[0]['content']}")
    if md_chunks:
        print(f"Markdown第一个块: {md_chunks[0]['content']}")
    if txt_chunks:
        print(f"TXT第一个块: {txt_chunks[0]['content']}")

    return {
        "csv": csv_chunks,
        "markdown": md_chunks,
        "txt": txt_chunks
    }


if __name__ == "__main__":
    results = example_usage()
    print("所有文件处理完成！")
